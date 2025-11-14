from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import json
from datetime import datetime
from urllib.parse import unquote
import unicodedata
from debate_system import DebateSystem, DebateSession
import random
import re # Added for regex validation

app = FastAPI(title="MLN Debate System API", version="1.0.0")

# Helper function to decode team_id
def decode_team_id(team_id: str) -> str:
    """Decode URL-encoded team_id to handle special characters"""
    return unquote(team_id)

def normalize_team_key(team_id: str) -> str:
    """Create a normalized key for team IDs to avoid Unicode/casing mismatches"""
    return unicodedata.normalize("NFKC", team_id).strip().lower()

def get_active_session(team_id: str):
    """Resolve and return the active session data by team identifier"""
    decoded_id = decode_team_id(team_id)
    session_key = normalize_team_key(decoded_id)
    session_data = active_sessions.get(session_key)
    if not session_data:
        print(f"❌ Session lookup failed for team_id='{decoded_id}'. Active sessions: {[data.get('team_id') for data in active_sessions.values()]}")
        raise HTTPException(status_code=404, detail="Session not found")
    return session_key, session_data

def get_any_session(team_id: str):
    """Resolve and return session data from either active or completed sessions"""
    decoded_id = decode_team_id(team_id)
    session_key = normalize_team_key(decoded_id)
    session_data = active_sessions.get(session_key)
    if session_data:
        return session_key, session_data

    # Fallback: iterate active sessions in case of unexpected key mismatch
    for key, active_data in active_sessions.items():
        if normalize_team_key(active_data.get("team_id", "")) == session_key:
            return key, active_data

    for session in completed_sessions:
        if normalize_team_key(session.get("team_id", "")) == session_key:
            return session_key, session
    for session in completed_sessions:
        if normalize_team_key(session.get("session_key", "")) == session_key:
            return session.get("session_key"), session
    print(f"❌ Session lookup failed (any) for team_id='{decoded_id}'.")
    raise HTTPException(status_code=404, detail="Session not found")

# CORS configuration - Fixed to prevent duplicate headers
# Updated for Vercel deployment - allows Vercel domains and localhost
import os
vercel_url = os.getenv("VERCEL_URL", "")
vercel_env = os.getenv("VERCEL_ENV", "")

# Build allowed origins list
allowed_origins = [
    "http://localhost:3000",
    "http://localhost:3001", 
    "http://127.0.0.1:3000",
    "http://127.0.0.1:3001",
    "http://mlndebate.io.vn",
    "https://mlndebate.io.vn",
    "http://www.mlndebate.io.vn",
    "https://www.mlndebate.io.vn",
]

blocked_patterns = [
    "tôi xin lỗi",
    "tôi không thể",
    "không thể trả lời",
    "as an ai",
    "i'm just an ai",
    "sorry",
    "cannot comply",
    "asdf",
    "ádfasd",
    "ấd",
    "ád",
]

# Add Vercel preview/deployment URLs dynamically
if vercel_url:
    allowed_origins.extend([
        f"https://{vercel_url}",
        f"http://{vercel_url}"
    ])

# For Vercel preview/development, we need to handle CORS differently
# Cannot use "*" with allow_credentials=True
allow_all_origins = vercel_env in ["development", "preview"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins if not allow_all_origins else ["*"],
    allow_credentials=not allow_all_origins,  # Cannot use credentials with wildcard
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
)

# CORS is now handled properly by CORSMiddleware

# Initialize debate system
try:
    debate_system = DebateSystem()
    print("✅ Debate system initialized successfully")
except Exception as e:
    print(f"⚠️ Warning: Could not initialize debate system: {e}")
    debate_system = None

# In-memory storage (replace with database in production)
active_sessions = {}
completed_sessions = []
session_counter = 0
class StartDebateRequest(BaseModel):
    course_code: str
    members: List[str]
    team_id: Optional[str] = None
    topic: Optional[str] = None

class SubmitArgumentsRequest(BaseModel):
    team_id: str
    arguments: List[str]

class SubmitQuestionRequest(BaseModel):
    team_id: str
    question: str

@app.get("/api/health")
async def health(): 
    return {
        "status": "healthy",
        "debate_system_available": debate_system is not None,
        "timestamp": datetime.now().isoformat()
    }

@app.post("/api/debate/start")
async def start_debate(request: StartDebateRequest):
    """Start a new debate session"""
    if not debate_system:
        raise HTTPException(status_code=503, detail="Debate system not available")
    
    # Use provided team_id or generate one
    if request.team_id and request.team_id.strip():
        team_id = request.team_id.strip()
    else:
        # Auto-generate if not provided
        global session_counter
        session_counter += 1
        team_id = f"TEAM{session_counter:03d}"

    session_key = normalize_team_key(team_id)
    # Check if normalized team_id already exists
    if session_key in active_sessions:
        raise HTTPException(status_code=400, detail=f"Team ID '{team_id}' already exists. Please choose a different one.")
    
    try:
        session = DebateSession(debate_system)
        topic = session.start_debate(request.course_code, request.members)
        
        # Randomly assign stance (agree/disagree)
        stance = random.choice(["agree", "disagree"])
        
        active_sessions[session_key] = {
            "session": session,
            "team_id": team_id,
            "session_key": session_key,
            "topic": topic,
            "members": request.members,
            "course_code": request.course_code,
            "status": "active",
            "current_phase": "Phase 1",
            "created_at": datetime.now().isoformat(),
            "turns_taken": 0,
            "stance": stance  # Add stance to session data
        }
        
        return {
            "success": True,
            "team_id": team_id,
            "topic": topic,
            "stance": stance,  # Return stance in response
            "message": f"Debate session started successfully. Your team will {'ĐỒNG TÌNH' if stance == 'agree' else 'PHẢN ĐỐI'} với chủ đề."
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to start debate: {str(e)}")

@app.post("/api/debate/{team_id}/arguments")
async def submit_arguments(team_id: str, request: SubmitArgumentsRequest):
    """Submit Phase 1 arguments"""
    try:
        _, session_data = get_active_session(team_id)
        session = session_data["session"]
        
        # 🔧 FIX: Store arguments in BOTH places for sync
        session.team_arguments = request.arguments  # ✅ For evaluation
        session_data["arguments"] = request.arguments  # ✅ For export
        session_data["current_phase"] = "Phase 2A"
        
        print(f"🔧 DEBUG: Stored team_arguments: {session.team_arguments}")
        
        # Generate AI questions based on arguments
        questions = debate_system.generate_questions(request.arguments, session_data["topic"])
        session_data["ai_questions"] = questions
        
        return {
            "success": True,
            "questions": questions,
            "message": "Arguments submitted successfully. AI has generated questions for Phase 2A."
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to submit arguments: {str(e)}")

@app.post("/api/debate/{team_id}/question")
async def submit_question(team_id: str, request: SubmitQuestionRequest):
    """Submit a question in Phase 2B"""
    try:
        _, session_data = get_active_session(team_id)
        
        # Generate Socratic response
        ai_response = debate_system.generate_socratic_answer(
            request.question, 
            session_data["topic"], 
            session_data.get("previous_context", "")
        )
        
        session_data["current_phase"] = "Phase 2B"
        session_data["turns_taken"] += 1
        
        return {
            "success": True,
            "ai_response": ai_response,
            "message": "Question submitted successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to submit question: {str(e)}")

@app.get("/api/debate/{team_id}/info")
async def get_debate_info(team_id: str):
    """Get debate session information"""
    try:
        _, session_data = get_active_session(team_id)
        return {
            "success": True,
            "team_id": session_data["team_id"],
            "topic": session_data["topic"],
            "members": session_data["members"],
            "course_code": session_data["course_code"],
            "status": session_data["status"],
            "current_phase": session_data["current_phase"],
            "turns_taken": session_data["turns_taken"],
            "created_at": session_data["created_at"],
            "arguments": session_data.get("arguments", []),
            "ai_questions": session_data.get("ai_questions", []),
            "stance": session_data.get("stance", "agree")  # Include stance in info
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get debate info: {str(e)}")

@app.get("/api/debate/{team_id}/turns")
async def get_debate_turns(team_id: str):
    """Get separated Phase 2 and Phase 3 turns"""
    try:
        _, session_data = get_active_session(team_id)
        session = session_data["session"]
        
        # Format Phase 2 turns (AI asks, Student answers)
        phase2_turns = []
        if hasattr(session, 'turns'):
            phase2_turns = [{
                "asker": turn.get("asker", "unknown"),
                "question": turn.get("question", ""),
                "answer": turn.get("answer"),
                "turn_number": idx + 1
            } for idx, turn in enumerate(session.turns)]
        
        # Format Phase 3 turns (Student asks, AI answers)
        phase3_turns = []
        if hasattr(session, 'phase3_turns'):
            phase3_turns = [{
                "asker": turn.get("asker", "unknown"),
                "question": turn.get("question", ""),
                "answer": turn.get("answer"),
                "turn_number": idx + 1
            } for idx, turn in enumerate(session.phase3_turns)]
        
        return {
            "success": True,
            "phase2_turns": phase2_turns,
            "phase3_turns": phase3_turns,
            "message": "Turns data retrieved successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get turns: {str(e)}")

class UpdatePhaseRequest(BaseModel):
    phase: str

class StanceRequest(BaseModel):
    stance: str

@app.post("/api/debate/{team_id}/stance")
async def set_stance(team_id: str, request: StanceRequest):
    """Set team stance (ĐỒNG TÌNH or PHẢN ĐỐI)"""
    try:
        _, session_data = get_active_session(team_id)
        session_data["stance"] = request.stance
        
        return {
            "success": True,
            "stance": request.stance,
            "message": f"Stance set to {request.stance}"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to set stance: {str(e)}")

@app.post("/api/debate/{team_id}/phase")
async def update_phase(team_id: str, request: UpdatePhaseRequest):
    """Update debate phase"""
    try:
        _, session_data = get_active_session(team_id)
        session_data["current_phase"] = request.phase
        
        return {
            "success": True,
            "current_phase": request.phase,
            "message": f"Phase updated to {request.phase}"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update phase: {str(e)}")

# Update Phase 1 to consider stance
@app.post("/api/debate/{team_id}/phase1")
async def get_ai_arguments_phase1(team_id: str):
    """Generate AI arguments for Phase 1"""
    if not debate_system:
        raise HTTPException(status_code=503, detail="Debate system not available")
    
    try:
        _, session_data = get_active_session(team_id)
        session = session_data["session"]
        topic = session_data["topic"]
        stance = session_data.get("stance", "agree")
        
        # Generate AI arguments opposing the team's stance
        ai_stance = "opposing" if stance == "agree" else "supporting"
        ai_arguments = debate_system.generate_arguments(topic, ai_stance)
        
        session.ai_arguments = ai_arguments
        session_data["ai_arguments"] = ai_arguments
        session_data["current_phase"] = "Phase 1.5"
        
        return {
            "success": True,
            "data": {
                "ai_arguments": ai_arguments,
                "topic": topic,
                "stance": stance,
                "message": "AI arguments generated successfully"
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate AI arguments: {str(e)}")

class Phase2Request(BaseModel):
    team_arguments: List[str]

@app.post("/api/debate/{team_id}/phase2")
async def get_ai_questions_phase2(team_id: str, request: Phase2Request):
    """Generate AI questions for Phase 2 and store student arguments"""
    if not debate_system:
        raise HTTPException(status_code=503, detail="Debate system not available")
    
    try:
        _, session_data = get_active_session(team_id)
        session = session_data["session"]
        topic = session_data["topic"]
        
        # 🔧 FIX: Store student arguments properly in BOTH places
        student_arguments = request.team_arguments
        session.team_arguments = student_arguments  # ✅ For evaluation
        session_data["arguments"] = student_arguments  # ✅ For export
        session_data["current_phase"] = "Phase 2"
        
        print(f"🔧 DEBUG Phase2: Stored team_arguments: {session.team_arguments}")
        
        if not student_arguments:
            # If no arguments, create some general ones to generate questions about
            student_arguments = [f"Ủng hộ quan điểm về chủ đề: {topic}"]
        
        # Generate AI questions challenging the students' position  
        if debate_system:
            ai_questions = debate_system.generate_questions(student_arguments, topic)
        else:
            ai_questions = ["Bạn có thể giải thích rõ hơn về quan điểm này không?"]
        
        # Store AI questions in session
        session_data["ai_questions"] = ai_questions
        
        # 🔧 FIX: Add the first AI question to session turns so it's tracked properly
        if ai_questions and len(ai_questions) > 0:
            session.add_turn("ai", ai_questions[0], None)
            print(f"🔧 DEBUG Phase2: Added first AI question to turns: {ai_questions[0][:50]}...")
        
        return {
            "success": True,
            "data": {
                "ai_questions": ai_questions,
                "topic": topic,
                "message": "AI questions generated successfully"
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate AI questions: {str(e)}")

@app.post("/api/debate/{team_id}/phase2/start")
async def start_phase2(team_id: str):
    """Start Phase 2 of the debate"""
    try:
        _, session_data = get_active_session(team_id)
        session_data["current_phase"] = "Phase 2"
        
        return {
            "success": True,
            "current_phase": "Phase 2",
            "message": "Phase 2 started successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to start Phase 2: {str(e)}")

class AIQuestionTurnRequest(BaseModel):
    answer: str
    asker: str
    question: str

@app.post("/api/debate/{team_id}/ai-question/turn")
async def ai_question_turn(team_id: str, request: AIQuestionTurnRequest):
    """Handle Phase 2: Student answers AI question and gets next AI question"""
    try:
        _, session_data = get_active_session(team_id)
        session = session_data["session"]
        
        # Validate student answer
        answer = request.answer.strip()
        if not answer or len(answer) < 10:
            raise HTTPException(status_code=400, detail="Câu trả lời quá ngắn (tối thiểu 10 ký tự)")
        
        # 🔧 DEBUG: Log before adding turn
        print(f"🔧 DEBUG ai_question_turn: About to add student answer")
        print(f"   Question: {request.question[:50]}{'...' if len(request.question) > 50 else ''}")
        print(f"   Answer: {answer[:50]}{'...' if len(answer) > 50 else ''}")
        print(f"   Current turns count: {len(session.turns)}")
        
        # 🔧 FIX: Add student answer turn with NO question (student only provides answers in Phase 2)
        session.add_turn("student", "", answer)
        
        # 🔧 DEBUG: Log after adding turn
        print(f"🔧 DEBUG ai_question_turn: After adding student answer, total turns: {len(session.turns)}")
        
        # 🔧 NEW FIX: Do NOT auto-generate next AI question
        # Let the frontend or user explicitly request the next question when ready
        # This prevents the "jumping to next question" issue
        print(f"🔧 BEHAVIOR: NOT auto-generating next AI question to prevent UI jumping")
        
        # Format turns for frontend
        formatted_turns = []
        for idx, turn in enumerate(session.turns):
            formatted_turn = {
                "asker": turn.get("asker", "unknown"),
                "question": turn.get("question", ""),
                "answer": turn.get("answer"),
                "turn_number": idx + 1
            }
            formatted_turns.append(formatted_turn)
        
        print(f"🔧 DEBUG: Phase 2 turns formatted. Total: {len(formatted_turns)}")
        
        return {
            "success": True,
            "turns": formatted_turns,
            "message": "Turn processed successfully"
        }
        
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process turn: {str(e)}")

@app.post("/api/debate/{team_id}/ai-question/generate")
async def generate_next_ai_question(team_id: str):
    """Generate next AI question for Phase 2 based on previous student answers"""
    try:
        _, session_data = get_active_session(team_id)
        session = session_data["session"]
        
        # Get the latest student answer to generate question from
        latest_student_answer = None
        for turn in reversed(session.turns):
            if turn.get('asker') == 'student' and turn.get('answer'):
                latest_student_answer = turn.get('answer')
                break
        
        if not latest_student_answer:
            raise HTTPException(status_code=400, detail="No student answer found to generate question from")
        
        # Generate next AI question
        next_ai_question = None
        
        # 🔧 RELAXED VALIDATION: More forgiving for test mode
        answer_clean = latest_student_answer.strip().lower()
        
        # More permissive validation - still block clear spam but allow test content
        severe_patterns = ['ádfasd', 'asdf', 'ấd', 'ád']  # Only severe nonsense patterns
        has_severe_nonsense = any(pattern in answer_clean for pattern in severe_patterns)
        is_extremely_short = len(latest_student_answer.strip()) < 5  # Very short only
        is_pure_numbers = bool(re.match(r'^[0-9\s]+$', latest_student_answer.strip()))  # Only pure numbers
        
        # SOFTER BLOCKING: Only block extremely obvious nonsense
        if has_severe_nonsense or is_extremely_short or is_pure_numbers:
            print(f"🚨 BLOCKED: Severe nonsense detected - using fallback only")
            next_ai_question = None
        else:
            # Only for CLEAN content - call AI system
            try:
                print(f"✅ CLEAN content detected - calling AI system")
                if session.debate_system:
                    ai_questions = session.debate_system.generate_questions(
                        [latest_student_answer],
                        session_data["topic"]
                    )
                    if ai_questions and len(ai_questions) > 0:
                        candidate_question = ai_questions[0].strip()
                        
                        # Final validation of AI response
                        is_clean_response = (
                            len(candidate_question) > 20 and 
                            '?' in candidate_question and
                            not any(pattern in candidate_question.lower() for pattern in blocked_patterns) and
                            not re.search(r'[0-9]{5,}', candidate_question) and
                            any(word in candidate_question.lower() for word in ['bạn', 'có', 'thể', 'như', 'nào', 'tại', 'sao', 'gì'])
                        )
                        
                        if is_clean_response:
                            next_ai_question = candidate_question
                            print(f"✅ AI response validated and accepted")
                        else:
                            print(f"🚨 AI response failed validation")
                            
            except Exception as e:
                print(f"Error in AI generation: {e}")
                
        # Add user guidance message if using fallback due to validation
        if not next_ai_question or "fallback" in next_ai_question.lower():
            print(f"ℹ️ INFO: Using fallback question. For better AI questions, provide substantive answers (avoid short/test responses)")
            
        # ALWAYS provide fallback if no valid AI question
        if not next_ai_question:
            topic_safe = re.sub(r'[0-9]{3,}', '', session_data["topic"])  # Remove long numbers from topic
            
            # 🔧 ANTI-REPETITION: Get previous questions to avoid duplicates
            previous_questions = set()
            for turn in session.turns:
                if turn.get('asker') == 'ai' and turn.get('question'):
                    previous_questions.add(turn.get('question').strip().lower())
            
            safe_fallbacks = [
                f"Bạn có thể phân tích sâu hơn về quan điểm của mình trong bối cảnh {topic_safe} không?",
                "Những bằng chứng nào có thể ủng hộ lập luận này?",
                "Bạn có thể so sánh với các quan điểm khác về vấn đề này không?",
                "Tác động thực tế của quan điểm này như thế nào?",
                "Những khía cạnh nào khác cần được xem xét?",
                "Bạn có thể giải thích rõ hơn về cơ sở lý thuyết không?",
                "Quan điểm này có những hạn chế gì cần thảo luận?",
                "Có những góc nhìn nào khác về vấn đề này?",
                "Bạn có thể đưa ra ví dụ cụ thể để minh họa không?",
                "Những thách thức chính của quan điểm này là gì?",
                "Làm sao để áp dụng quan điểm này vào thực tế?",
                "Có những nghiên cứu nào ủng hộ quan điểm này?",
                "Bạn có thể phân tích ưu và nhược điểm không?"
            ]
            
            # 🔧 SMART SELECTION: Choose question not used before
            available_questions = [q for q in safe_fallbacks if q.strip().lower() not in previous_questions]
            if available_questions:
                next_ai_question = random.choice(available_questions)
                print(f"🎯 Using new fallback: {next_ai_question[:60]}...")
            else:
                # If all questions used, add variety with current turn number
                turn_num = len(session.turns) + 1
                next_ai_question = f"Ở góc độ thứ {turn_num}, bạn có thể làm rõ thêm về vấn đề này không?"
                print(f"🔄 Using numbered fallback: {next_ai_question[:60]}...")
            
        # Add the question to session
        session.add_turn("ai", next_ai_question, None)
        
        # Format turns for frontend
        formatted_turns = []
        for idx, turn in enumerate(session.turns):
            formatted_turn = {
                "asker": turn.get("asker", "unknown"),
                "question": turn.get("question", ""),
                "answer": turn.get("answer"),
                "turn_number": idx + 1
            }
            formatted_turns.append(formatted_turn)
        
        return {
            "success": True,
            "turns": formatted_turns,
            "new_question": next_ai_question,
            "message": "Next AI question generated successfully"
        }
        
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate next question: {str(e)}")

class StudentQuestionTurnRequest(BaseModel):
    asker: str
    question: str
    answer: Optional[str] = None

@app.post("/api/debate/{team_id}/student-question/turn")
async def student_question_turn(team_id: str, request: StudentQuestionTurnRequest):
    """Handle Phase 3: Student asks question and gets AI answer"""
    try:
        _, session_data = get_active_session(team_id)
        session = session_data["session"]
        
        cleaned_question = request.question.strip()
        if len(cleaned_question) < 12 or '?' not in cleaned_question or not re.search(r'[a-zA-ZÀ-ỹ]', cleaned_question):
            raise HTTPException(
                status_code=400,
                detail="Câu hỏi chưa đủ rõ ràng. Vui lòng đặt lại với nội dung cụ thể và có dấu chấm hỏi."
            )
        
        # 🔧 FIX: Use session.add_phase3_turn() for Phase 3 data
        # First add student question
        session.add_phase3_turn("student", cleaned_question, None)
        
        # Generate AI answer using Socratic method
        try:
            if session.debate_system:
                ai_answer = session.debate_system.generate_socratic_answer(
                    student_question=request.question.strip(),
                    topic=session_data["topic"],
                    previous_context=""
                )
            else:
                ai_answer = "Hệ thống AI tạm thời không khả dụng."
            
            if ai_answer:
                # Add AI answer using session.add_phase3_turn() - AI answers don't have questions, only answers
                session.add_phase3_turn("ai", None, ai_answer)
            
        except Exception as e:
            print(f"Error generating AI answer: {e}")
            # If AI generation fails, add default response
            session.add_phase3_turn(
                "ai",
                None,
                "Tôi không chắc mình hiểu câu hỏi của bạn. Bạn có thể diễn đạt rõ hơn hoặc nêu cụ thể điều muốn hỏi không?"
            )
        
        print(f"🔧 DEBUG: Phase 3 turns added. Total phase3_turns: {len(session.phase3_turns)}")
        
        return {
            "success": True,
            "turns": session.phase3_turns,  # Return Phase 3 turns list
            "message": "Question processed successfully"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process question: {str(e)}")

@app.delete("/api/debate/{team_id}/end")
async def end_debate(team_id: str):
    """End/delete a debate session"""
    try:
        session_key, session_data = get_active_session(team_id)
        
        # Move to completed sessions as "ended"
        completed_session = {
            **session_data,
            "status": "ended",
            "completed_at": datetime.now().isoformat(),
            "end_reason": "manual_end"
        }
        completed_sessions.append(completed_session)
        del active_sessions[session_key]
        
        return {
            "success": True,
            "message": f"Debate {team_id} ended successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to end debate: {str(e)}")

@app.post("/api/debate/{team_id}/complete")
async def complete_debate(team_id: str):
    """Complete a debate session after Phase 5 evaluation"""
    try:
        session_key, session_data = get_active_session(team_id)
        
        # Check if evaluation exists (Phase 5 completed)
        if "evaluation" not in session_data:
            raise HTTPException(status_code=400, detail="Please complete Phase 5 evaluation first")
        
        # Move to completed sessions
        completed_session = {
            **session_data,
            "status": "completed",
            "completed_at": datetime.now().isoformat()
        }
        completed_sessions.append(completed_session)
        del active_sessions[session_key]
        
        return {
            "success": True,
            "evaluation": session_data["evaluation"],
            "message": "Debate session completed and archived successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to complete debate: {str(e)}")

@app.get("/api/admin/sessions")  
async def get_sessions(): 
    """Get all active and completed sessions"""
    try:
        # Convert active sessions to API format
        active = []
        for session_key, session_data in active_sessions.items():
            active.append({
                "team_id": session_data.get("team_id", session_key),
                "topic": session_data["topic"],
                "status": session_data["status"],
                "current_phase": session_data["current_phase"],
                "members": session_data["members"],
                "turns_taken": session_data["turns_taken"],
                "course_code": session_data.get("course_code", "")
            })
        
        # Convert completed sessions to API format
        completed = []
        for session_data in completed_sessions[-10:]:  # Last 10 completed
            # For sessions that were force-ended without evaluation, create basic evaluation
            evaluation = session_data.get("evaluation")
            if not evaluation and session_data.get("status") == "ended":
                # Import criteria from debate_system 
                from debate_system import DEBATE_CRITERIA
                
                scores = {}
                for phase_key, criteria_list in DEBATE_CRITERIA.items():
                    scores[phase_key] = {criterion['id']: 0 for criterion in criteria_list}
                
                evaluation = {
                    "total_score": 0,
                    "scores": scores,
                    "feedback": "Session was ended manually before completion. No detailed evaluation available."
                }
                session_data["evaluation"] = evaluation  # Save it back
            
            completed.append({
                "team_id": session_data["team_id"],
                "topic": session_data["topic"],
                "status": session_data["status"],
                "completed_at": session_data["completed_at"],
                "members": session_data["members"],
                "score": evaluation.get("total_score", 0) if evaluation else 0,
                "evaluation": evaluation  # Include full evaluation data
            })
        
        # Import criteria from debate_system 
        from debate_system import DEBATE_CRITERIA
        
        return {
            "active": active,
            "completed": completed,
            "criteria": DEBATE_CRITERIA
        }
    except Exception as e:
        print(f"Error in get_sessions: {e}")
        # Return mock data as fallback
        # Import criteria from debate_system 
        from debate_system import DEBATE_CRITERIA
        
        return {
            "active": [],
            "completed": [],
            "criteria": DEBATE_CRITERIA
        }

@app.get("/api/admin/leaderboard")
async def get_leaderboard(): 
    """Get leaderboard from completed sessions"""
    try:
        # Calculate leaderboard from completed sessions
        leaderboard_data = []
        for i, session in enumerate(completed_sessions):
            evaluation = session.get("evaluation", {})
            scores = evaluation.get("scores", {})
            
            # Calculate total score
            total_score = 0
            phase_scores = {}
            for phase, phase_scores_dict in scores.items():
                if isinstance(phase_scores_dict, dict):
                    phase_total = sum(phase_scores_dict.values())
                    phase_scores[phase] = phase_total
                    total_score += phase_total
            
            leaderboard_data.append({
                "position": i + 1,
                "team_id": session["team_id"],
                "course_code": session.get("course_code", "MLN111"),
                "topic": session["topic"],
                "members": session["members"],
                "total_score": total_score,
                "max_score": 125,
                "percentage": min(100, (total_score / 125) * 100) if total_score > 0 else 0,
                "rank_level": "Gold Level" if total_score > 80 else "Silver Level",
                "phase_scores": phase_scores,
                "completed_at": session.get("completed_at", "2024-01-01T00:00:00")
            })
        
        # Sort by total score
        leaderboard_data.sort(key=lambda x: x["total_score"], reverse=True)
        
        # Update positions
        for i, entry in enumerate(leaderboard_data):
            entry["position"] = i + 1
        
        return {
            "leaderboard": leaderboard_data[:20],  # Top 20
            "statistics": {
                "total_teams": len(completed_sessions),
                "average_score": sum(entry["total_score"] for entry in leaderboard_data) / len(leaderboard_data) if leaderboard_data else 0,
                "highest_score": max((entry["total_score"] for entry in leaderboard_data), default=0),
                "rank_distribution": {
                    "Platinum": len([e for e in leaderboard_data if e["total_score"] > 95]),
                    "Gold": len([e for e in leaderboard_data if 80 <= e["total_score"] <= 95]),
                    "Silver": len([e for e in leaderboard_data if 60 <= e["total_score"] < 80]),
                    "Bronze": len([e for e in leaderboard_data if e["total_score"] < 60])
                }
            }
        }
    except Exception as e:
        print(f"Error in get_leaderboard: {e}")
        # Return mock data as fallback
        return {
            "leaderboard": [],
            "statistics": {
                "total_teams": 0,
                "average_score": 0,
                "highest_score": 0,
                "rank_distribution": {"Platinum": 0, "Gold": 0, "Silver": 0, "Bronze": 0}
            }
        }

@app.get("/api/admin/live-scoring")
async def get_live_scoring(): 
    """Get live scoring data"""
    try:
        live_data = []
        for session_key, session_data in active_sessions.items():
            live_data.append({
                "team_id": session_data.get("team_id", session_key),
                "topic": session_data["topic"],
                "status": "in_progress",
                "current_phase": session_data["current_phase"],
                "members": session_data["members"],
                "progress": min(100, session_data["turns_taken"] * 20),  # Rough progress calculation
                "elapsed_time": "00:15:30"  # Could calculate actual time
            })
        
        return {
            "live_scoring": live_data,
            "statistics": {
                "active_debates": len(active_sessions),
                "total_participants": sum(len(session["members"]) for session in active_sessions.values()),
                "average_progress": sum(min(100, session["turns_taken"] * 20) for session in active_sessions.values()) / len(active_sessions) if active_sessions else 0
            }
        }
    except Exception as e:
        print(f"Error in get_live_scoring: {e}")
        return {
            "live_scoring": [],
            "statistics": {
                "active_debates": 0,
                "total_participants": 0,
                "average_progress": 0
            }
        }

@app.post("/api/debate/{team_id}/phase4/conclusion")
async def submit_conclusion(team_id: str, request: SubmitArgumentsRequest):
    """Phase 4 Step 1: Submit student final conclusion - why they should win"""
    try:
        session_key, session_data = get_active_session(team_id)
        session = session_data["session"]
        
        # Check if conclusion already exists
        if "conclusion" in session_data and session_data["conclusion"]:
            return {
                "success": True,
                "conclusion": session_data["conclusion"],
                "message": "Student conclusion already submitted."
            }
        
        # Filter out empty arguments and ensure we have valid content
        valid_arguments = [arg.strip() for arg in request.arguments if arg and arg.strip()]
        
        # If no valid arguments, use default conclusion
        if not valid_arguments:
            valid_arguments = [
                "Nhóm chúng tôi có luận điểm vững chắc được trình bày trong các giai đoạn trước.",
                "Các câu trả lời của chúng tôi thể hiện sự hiểu biết sâu sắc về chủ đề.",
                "Chúng tôi đã phản biện hiệu quả các luận điểm của AI."
            ]
        
        # Store student conclusion arguments
        session_data["conclusion"] = valid_arguments
        session.conclusion = valid_arguments  # Sync with DebateSession
        session_data["current_phase"] = "Phase 4 - Student Conclusion"
        
        return {
            "success": True,
            "conclusion": valid_arguments,
            "message": "Phase 4 Step 1 completed: Student conclusion submitted. Now AI will generate counter-arguments."
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to submit conclusion: {str(e)}")

@app.post("/api/debate/{team_id}/phase5/evaluate")
async def evaluate_debate_phase5(team_id: str):
    """Phase 5: Final evaluation and scoring"""
    try:
        session_key, session_data = get_active_session(team_id)
        session = session_data["session"]
        
        # Generate comprehensive evaluation
        evaluation = session.evaluate_debate()
        
        # Update session data
        session_data["current_phase"] = "Phase 5"
        session_data["evaluation"] = evaluation
        
        return {
            "success": True,
            "evaluation": evaluation,
            "message": "Phase 5 evaluation completed successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to evaluate debate: {str(e)}")

@app.get("/api/debate/{team_id}/phase4/info")
async def get_phase4_info(team_id: str):
    """Get Phase 4 conclusion information"""
    try:
        session_key, session_data = get_active_session(team_id)
        return {
            "success": True,
            "team_id": session_data["team_id"],
            "topic": session_data["topic"],
            "current_phase": session_data["current_phase"],
            "conclusion": session_data.get("conclusion", []),
            "ai_counter_arguments": session_data.get("ai_counter_arguments", []),
            "arguments": session_data.get("arguments", []),
            "message": "Phase 4 information retrieved successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get phase 4 info: {str(e)}")

@app.post("/api/debate/{team_id}/phase4/evaluate")
async def evaluate_phase4(team_id: str):
    """Phase 4 Step 3: Mark Phase 4 as completed after AI counter-conclusion"""
    try:
        session_key, session_data = get_active_session(team_id)
        
        # Check if already completed
        if session_data.get("current_phase") == "Phase 4 Completed":
            return {
                "success": True,
                "current_phase": "Phase 4 Completed",
                "message": "Phase 4 already completed."
            }
        
        # Check if both student conclusion and AI counter-arguments exist
        if "conclusion" not in session_data or not session_data["conclusion"]:
            raise HTTPException(status_code=400, detail="Please submit student conclusion first")
        
        if "ai_counter_arguments" not in session_data or not session_data["ai_counter_arguments"]:
            raise HTTPException(status_code=400, detail="Please generate AI counter-arguments first")
        
        # Mark Phase 4 as completed
        session_data["current_phase"] = "Phase 4 Completed"
        
        return {
            "success": True,
            "current_phase": "Phase 4 Completed",
            "message": "Phase 4 evaluation completed. Ready for Phase 5."
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to evaluate Phase 4: {str(e)}")

@app.post("/api/debate/{team_id}/phase4/ai-conclusion")
async def generate_ai_conclusion(team_id: str):
    """Phase 4 Step 2: Generate AI counter-conclusion - why AI should win"""
    if not debate_system:
        raise HTTPException(status_code=503, detail="Debate system not available")
    
    try:
        session_key, session_data = get_active_session(team_id)
        session = session_data["session"]
        
        # Check if AI counter-arguments already exist
        if "ai_counter_arguments" in session_data and session_data["ai_counter_arguments"]:
            return {
                "success": True,
                "ai_counter_arguments": session_data["ai_counter_arguments"],
                "message": "AI counter-arguments already generated."
            }
        
        # Check if student has submitted conclusion
        if "conclusion" not in session_data:
            raise HTTPException(status_code=400, detail="Please submit student conclusion first")
        
        topic = session_data["topic"]
        student_conclusion = session_data["conclusion"]
        
        # Generate AI counter-conclusion (why AI should win)
        ai_counter_arguments = debate_system.generate_arguments(
            f"Tại sao AI nên thắng trong cuộc tranh luận về chủ đề '{topic}'. Phản bác lại các luận điểm tổng kết của sinh viên: " + "; ".join(student_conclusion), 
            "opposing"
        )
        
        # Store AI counter-arguments
        session_data["ai_counter_arguments"] = ai_counter_arguments
        session.ai_counter_arguments = ai_counter_arguments  # Sync with DebateSession
        session_data["current_phase"] = "Phase 4 - AI Conclusion"
        
        return {
            "success": True,
            "ai_counter_arguments": ai_counter_arguments,
            "message": "Phase 4 Step 2 completed: AI counter-conclusion generated. Ready for Phase 5 evaluation."
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate AI conclusion: {str(e)}")

@app.get("/api/debate/{team_id}/export_docx")
async def export_debate_report(team_id: str):
    """Export debate report as DOCX file with complete debate history"""
    decoded_id = decode_team_id(team_id)

    _, session_data = get_any_session(team_id)
    session_obj = session_data.get("session")
    team_id_display = session_data.get("team_id", decoded_id)
    
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        
        # Create document
        doc = Document()
        doc.add_heading('MLN Debate System - Báo Cáo Chi Tiết & Lịch Sử Tranh Luận', 0)
        
        # Team information
        doc.add_heading('📋 Thông Tin Nhóm', level=1)
        doc.add_paragraph(f"Team ID: {team_id_display}")
        doc.add_paragraph(f"Chủ đề: {session_data.get('topic', 'N/A')}")
        doc.add_paragraph(f"Thành viên: {', '.join(session_data.get('members', []))}")
        doc.add_paragraph(f"Mã học phần: {session_data.get('course_code', 'N/A')}")
        doc.add_paragraph(f"Thời gian tạo: {session_data.get('created_at', 'N/A')}")
        doc.add_paragraph(f"Trạng thái: {session_data.get('status', 'N/A')}")
        if session_data.get('completed_at'):
            doc.add_paragraph(f"Thời gian hoàn thành: {session_data.get('completed_at', 'N/A')}")
        
        # DEBATE HISTORY SECTION
        doc.add_heading('🎯 Lịch Sử Tranh Luận Chi Tiết', level=1)
        
        # Phase 1: Initial Arguments
        doc.add_heading('Phase 1: Luận Điểm Ban Đầu', level=2)
        
        # Team Arguments
        team_arguments = session_data.get('arguments', [])
        if session_obj and hasattr(session_obj, 'team_arguments'):
            team_arguments = session_obj.team_arguments
        
        if team_arguments:
            doc.add_heading('💭 Luận điểm của Team:', level=3)
            for i, arg in enumerate(team_arguments, 1):
                doc.add_paragraph(f"{i}. {arg}", style='List Number')
        else:
            doc.add_paragraph("Chưa có luận điểm từ team.")
        
        # AI Arguments  
        ai_arguments = session_data.get('ai_arguments', [])
        if session_obj and hasattr(session_obj, 'ai_arguments'):
            ai_arguments = session_obj.ai_arguments
            
        if ai_arguments:
            doc.add_heading('🤖 Luận điểm của AI:', level=3)
            for i, arg in enumerate(ai_arguments, 1):
                doc.add_paragraph(f"{i}. {arg}", style='List Number')
        else:
            doc.add_paragraph("Chưa có luận điểm từ AI.")
        
        # Phase 2: AI Questions & Team Responses
        doc.add_heading('Phase 2: AI Chất Vấn Team', level=2)
        
        ai_questions = session_data.get('ai_questions', [])
        if session_obj and hasattr(session_obj, 'questions'):
            ai_questions = session_obj.questions
            
        if ai_questions:
            doc.add_heading('❓ Câu hỏi của AI:', level=3)
            for i, question in enumerate(ai_questions, 1):
                doc.add_paragraph(f"Q{i}: {question}", style='Intense Quote')
                
        # Phase 2 Dialog: AI Questions & Student Responses
        if session_obj and hasattr(session_obj, 'turns') and session_obj.turns:
            doc.add_heading('🔄 Cuộc hội thoại Phase 2 (AI chất vấn Team):', level=3)
            
            # 🔧 FIXED: Phase 2 = ALL turns from session.turns
            # The session.turns array contains ONLY Phase 2 data
            # Phase 3 data is stored separately in session.phase3_turns
            
            ai_questions = []
            student_answers = []
            
            # Process ALL turns from session.turns (these are ALL Phase 2)
            for turn in session_obj.turns:
                if turn.get('asker') == 'ai' and turn.get('question'):
                    ai_questions.append(turn.get('question', ''))
                elif turn.get('asker') == 'student' and turn.get('answer'):
                    student_answers.append(turn.get('answer', ''))
            
            # Create pairs by matching questions with answers sequentially
            turn_pairs = []
            max_pairs = max(len(ai_questions), len(student_answers))
            
            for i in range(max_pairs):
                ai_question = ai_questions[i] if i < len(ai_questions) else ''
                student_answer = student_answers[i] if i < len(student_answers) else ''
                
                if ai_question or student_answer:  # Only add if there's content
                    turn_pairs.append({
                        'ai_question': ai_question,
                        'student_answer': student_answer
                    })
            
            # Display pairs nicely
            for i, pair in enumerate(turn_pairs, 1):
                doc.add_paragraph(f"Lượt {i}:", style='Heading 4')
                doc.add_paragraph(f"🤖 AI hỏi: {pair.get('ai_question', '')}")
                if pair.get('student_answer'):
                    doc.add_paragraph(f"👥 Team trả lời: {pair.get('student_answer', '')}", style='Intense Quote')
                else:
                    doc.add_paragraph(f"👥 Team trả lời: (Chưa trả lời)")
                doc.add_paragraph()  # Empty line
        
        # Phase 3: Team Questions & AI Responses  
        doc.add_heading('Phase 3: Team Chất Vấn AI', level=2)
        
        # 🔧 FIXED: Use dedicated phase3_turns array - no complex logic needed
        if session_obj and hasattr(session_obj, 'phase3_turns') and session_obj.phase3_turns:
            doc.add_heading('🔄 Lượt hỏi đáp Phase 3:', level=3)
            
            # Group Phase 3 turns by pairs (Student question + AI answer)
            phase3_pairs = []
            current_phase3_pair = {}
            
            for turn in session_obj.phase3_turns:
                if turn.get('asker') == 'student' and turn.get('question'):
                    if current_phase3_pair:  # Save previous pair
                        phase3_pairs.append(current_phase3_pair)
                    current_phase3_pair = {'student_question': turn.get('question', ''), 'ai_answer': ''}
                elif turn.get('asker') == 'ai' and turn.get('answer'):
                    if current_phase3_pair:
                        current_phase3_pair['ai_answer'] = turn.get('answer', '')
            
            if current_phase3_pair:  # Add last pair
                phase3_pairs.append(current_phase3_pair)
            
            # Display Phase 3 pairs
            for i, pair in enumerate(phase3_pairs, 1):
                doc.add_paragraph(f"Lượt {i}:", style='Heading 4')
                doc.add_paragraph(f"👥 Team hỏi: {pair.get('student_question', '')}")
                if pair.get('ai_answer'):
                    doc.add_paragraph(f"🤖 AI trả lời: {pair.get('ai_answer', '')}", style='Intense Quote')
                else:
                    doc.add_paragraph(f"🤖 AI trả lời: (Đang chờ AI trả lời...)")
                doc.add_paragraph()  # Empty line
        else:
            doc.add_paragraph("(Chưa có lượt hỏi đáp nào trong Phase 3)")
        
        # Phase 4: Final Conclusions
        doc.add_heading('Phase 4: Kết Luận Cuối Cùng', level=2)
        
        # Student Conclusion
        conclusion = session_data.get('conclusion', [])
        if session_obj and hasattr(session_obj, 'conclusion'):
            conclusion = session_obj.conclusion
            
        if conclusion:
            doc.add_heading('🎯 Kết luận của Team (Tại sao team nên thắng):', level=3)
            for i, conc in enumerate(conclusion, 1):
                doc.add_paragraph(f"{i}. {conc}", style='List Number')
        else:
            doc.add_paragraph("Chưa có kết luận từ team.")
        
        # AI Counter-arguments
        ai_counter = session_data.get('ai_counter_arguments', [])
        if session_obj and hasattr(session_obj, 'ai_counter_arguments'):
            ai_counter = session_obj.ai_counter_arguments
            
        if ai_counter:
            doc.add_heading('🤖 Phản bác của AI (Tại sao AI nên thắng):', level=3)
            for i, counter in enumerate(ai_counter, 1):
                doc.add_paragraph(f"{i}. {counter}", style='List Number')
        else:
            doc.add_paragraph("Chưa có phản bác từ AI.")
        
        # 🚫 REMOVED: Chat History section completely to prevent data mixing
        # All conversation data is already displayed in proper Phase 2 and Phase 3 sections above
        
        # Evaluation scores
        if session_data.get('evaluation'):
            evaluation = session_data['evaluation']
            doc.add_heading('📊 Kết Quả Chấm Điểm', level=1)
            
            total_score = 0
            total_max_score = 100
            
            for phase_key in ['phase1', 'phase2', 'phase3', 'phase4']:
                if phase_key in evaluation.get('scores', {}):
                    phase_scores = evaluation['scores'][phase_key]
                    phase_name = {
                        'phase1': 'Giai đoạn 1: Luận điểm ban đầu',
                        'phase2': 'Giai đoạn 2: AI chất vấn SV', 
                        'phase3': 'Giai đoạn 3: SV chất vấn AI',
                        'phase4': 'Giai đoạn 4: Tổng kết luận điểm'
                    }.get(phase_key, phase_key)

                    doc.add_heading(phase_name, level=2)
                    
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Tiêu chí'
                    hdr_cells[1].text = 'Điểm'
                    hdr_cells[2].text = 'Tối đa'

                    phase_total = 0
                    for criterion_id, score in phase_scores.items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = f"Tiêu chí {criterion_id}"
                        row_cells[1].text = str(score)
                        row_cells[2].text = "5" if criterion_id.endswith(('.3', '.4', '.5', '.6')) else "6"
                        phase_total += int(score) if score else 0
                    
                    # Add phase total row
                    total_row = table.add_row().cells
                    total_row[0].text = 'Tổng điểm giai đoạn'
                    total_row[1].text = str(phase_total)
                    total_row[2].text = "25"
                    
                    total_score += phase_total

            # Grand total
            doc.add_heading('🏆 Tổng Điểm', level=2)
            doc.add_paragraph(f"Tổng điểm: {total_score} / {total_max_score}")
            doc.add_paragraph(f"Tỷ lệ: {(total_score/total_max_score*100):.1f}%")
            
            # Feedback
            if evaluation.get('feedback'):
                doc.add_heading('💭 Nhận Xét Từ AI', level=2)
                doc.add_paragraph(evaluation['feedback'])
        
        # Summary Statistics
        doc.add_heading('📈 Thống Kê Tổng Quan', level=1)
        total_turns = 0
        if session_obj and hasattr(session_obj, 'turns'):
            total_turns = len(session_obj.turns)
        if session_obj and hasattr(session_obj, 'phase3_turns'):
            total_turns += len(session_obj.phase3_turns)
            
        doc.add_paragraph(f"Tổng số lượt hỏi đáp: {total_turns}")
        doc.add_paragraph(f"Số luận điểm team: {len(team_arguments)}")
        doc.add_paragraph(f"Số luận điểm AI: {len(ai_arguments)}")
        doc.add_paragraph(f"Giai đoạn hiện tại: {session_data.get('current_phase', 'N/A')}")
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph("📋 Báo cáo được tạo bởi MLN Debate System")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para = doc.add_paragraph("🌟 Hệ thống hỗ trợ tranh luận học thuật với AI")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=debate_full_report_{team_id}.docx"}
        )
        
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library not installed")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate report: {str(e)}")

@app.delete("/api/admin/history/{team_id}")
async def delete_session_history(team_id: str):
    """Delete a session from completed history"""
    decoded_id = decode_team_id(team_id)
    normalized_key = normalize_team_key(decoded_id)
    
    try:
        # Find and remove from completed sessions
        for i, session in enumerate(completed_sessions):
            if normalize_team_key(session.get("team_id", "")) == normalized_key:
                removed_session = completed_sessions.pop(i)
                return {
                    "success": True,
                    "message": f"Session {removed_session.get('team_id', decoded_id)} deleted from history successfully",
                    "deleted_session": {
                        "team_id": removed_session["team_id"],
                        "topic": removed_session.get("topic", "N/A"),
                        "status": removed_session.get("status", "unknown")
                    }
                }
        
        # If not found in completed, check if it's in active sessions
        if normalized_key in active_sessions:
            raise HTTPException(status_code=400, detail="Cannot delete active session. Please end it first.")
        
        # Session not found
        raise HTTPException(status_code=404, detail="Session not found in history")
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete session: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)
